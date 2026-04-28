# -*- coding: utf-8 -*-
"""
创建测试用Word和Excel文件
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


def create_test_word():
    """创建测试Word文件"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("请先安装 python-docx: pip install python-docx")
        return
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('Work Instruction - Assembly Process', level=1)
    
    # 添加段落
    doc.add_paragraph('1. Install the mounting bracket onto the frame using four M8 bolts.')
    doc.add_paragraph('2. Torque the bolts to 25 Nm using a torque wrench.')
    doc.add_paragraph('3. Connect the wiring harness to the control module.')
    
    # 添加警告
    warning = doc.add_paragraph()
    warning_run = warning.add_run('Warning: ')
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    warning.add_run('Disconnect the power supply before performing maintenance.')
    
    # 添加表格
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Step'
    hdr_cells[1].text = 'Description'
    
    # 数据行
    row1 = table.rows[1].cells
    row1[0].text = '1'
    row1[1].text = 'Install bracket'
    
    row2 = table.rows[2].cells
    row2[0].text = '2'
    row2[1].text = 'Tighten bolts'
    
    # 保存
    output_path = os.path.join(os.path.dirname(__file__), "sample.docx")
    doc.save(output_path)
    print(f"已创建测试Word文件: {output_path}")


def create_test_excel():
    """创建测试Excel文件"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
    except ImportError:
        print("请先安装 openpyxl: pip install openpyxl")
        return
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Instructions"
    
    # 表头
    ws['A1'] = 'Step'
    ws['B1'] = 'English'
    ws['C1'] = 'Chinese'
    
    # 设置表头格式
    for cell in ['A1', 'B1', 'C1']:
        ws[cell].font = Font(bold=True)
        ws[cell].fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
    
    # 数据
    data = [
        [1, 'Install the mounting bracket', ''],
        [2, 'Torque the bolts to 25 Nm', ''],
        [3, 'Connect the wiring harness', ''],
        [4, 'Verify all connections', ''],
    ]
    
    for i, row in enumerate(data, start=2):
        ws[f'A{i}'] = row[0]
        ws[f'B{i}'] = row[1]
        ws[f'C{i}'] = row[2]
    
    # 调整列宽
    ws.column_dimensions['A'].width = 10
    ws.column_dimensions['B'].width = 40
    ws.column_dimensions['C'].width = 40
    
    # 保存
    output_path = os.path.join(os.path.dirname(__file__), "sample.xlsx")
    wb.save(output_path)
    print(f"已创建测试Excel文件: {output_path}")


if __name__ == "__main__":
    print("创建测试文件...")
    create_test_word()
    create_test_excel()
    print("完成!")
